"""
Renders a generated Markdown document to .docx, following the project's
Word conventions: plain black text, no Word "Heading" styles (headers are
bold plain paragraphs instead), section titles in bold only.

Handles exactly the Markdown constructs scripts/make_week2_handback.py
produces: #/##/### headers, pipe tables, ```-fenced code blocks, **bold**,
`inline code`, and plain paragraphs -- not a general-purpose Markdown
renderer.

Usage: PYTHONPATH=. python scripts/render_docx.py <input.md> <output.docx>
"""
import re
import sys

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLACK = RGBColor(0, 0, 0)
BODY_FONT = "Calibri"
CODE_FONT = "Consolas"


def _set_plain_run(run, bold=False, mono=False, size=11):
    run.font.color.rgb = BLACK
    run.font.name = CODE_FONT if mono else BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold


def add_inline_formatted(paragraph, text: str):
    """Splits text on **bold** and `code` spans and adds runs accordingly."""
    token_re = re.compile(r"(\*\*.+?\*\*|`.+?`)")
    for token in token_re.split(text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            _set_plain_run(run, bold=True)
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            _set_plain_run(run, mono=True, size=10)
        else:
            run = paragraph.add_run(token)
            _set_plain_run(run)


def add_heading(doc, text: str, level: int):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    size = {1: 16, 2: 14, 3: 12}.get(level, 11)
    run = p.add_run(text)
    _set_plain_run(run, bold=True, size=size)


def add_code_block(doc, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Pt(18)
    for i, line in enumerate(code.split("\n")):
        run = p.add_run(line)
        _set_plain_run(run, mono=True, size=9)
        if i < len(code.split("\n")) - 1:
            p.add_run().add_break()


def add_table(doc, header: list, rows: list):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        p = hdr_cells[i].paragraphs[0]
        add_inline_formatted(p, h)
        for run in p.runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                # cells[i].paragraphs[0] already exists (empty) -- reuse it
                # rather than add_paragraph(), which would leave a stray blank one.
                add_inline_formatted(cells[i].paragraphs[0], val)


def parse_table_block(lines: list) -> tuple:
    header = [c.strip() for c in lines[0].strip().strip("|").split("|")]
    data_rows = []
    for line in lines[2:]:
        if not line.strip().startswith("|"):
            break
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        data_rows.append(cells)
    return header, data_rows


def render(md_path: str, docx_path: str):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]

        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, "\n".join(code_lines))
            i += 1  # skip closing ```
            continue

        if line.startswith("#"):
            level = len(line) - len(line.lstrip("#"))
            text = line.lstrip("#").strip()
            add_heading(doc, text, level)
            i += 1
            continue

        if line.strip().startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header, rows = parse_table_block(table_lines)
            add_table(doc, header, rows)
            doc.add_paragraph()
            continue

        if line.strip() == "":
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_formatted(p, line.strip()[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_formatted(p, line)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: PYTHONPATH=. python scripts/render_docx.py <input.md> <output.docx>")
        sys.exit(1)
    render(sys.argv[1], sys.argv[2])
