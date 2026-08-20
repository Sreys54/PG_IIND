"""
Week 3, Entregable 10 (extended Week 4, Entregable 13): appends a new
section per week to the user's personal Progress_Log_Thesis_Project.docx
WITHOUT rewriting anything already in it.

That file lives OUTSIDE this git repository, one directory above the repo
root (../Progress_Log_Thesis_Project.docx relative to PG_IIND/) -- it is
the user's own document, not a build artifact of this repo, so it is not
committed or regenerated from scratch the way the other hand-back
documents are. This script only appends.

Note found while building this script: the file's Weekly Progress Log
section (part 2) contains a "2.1. Week 1" subsection but NO "2.2. Week 2"
subsection -- Week 2's own instructions apparently asked for this file to
be extended too, but that never happened. This script adds Week 3 as
"2.2." (continuing the document's own existing "2.<n>. Week <n>" numbering
scheme, not the "3.x" numbering guessed at in the Week 3 task brief, since
the document's own established convention takes precedence) -- it does
NOT attempt to backfill a Week 2 section, since fabricating that content
now, after the fact, risks misrepresenting what was actually reported
in real time. Flagged for the user rather than silently patched over.

Usage: PYTHONPATH=. python scripts/extend_progress_log.py
"""
import os

from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx import Document

PROGRESS_LOG_PATH = "../Progress_Log_Thesis_Project.docx"

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"
SIZE = 12


def _set_run(run, bold=False):
    run.font.name = FONT
    run.font.size = Pt(SIZE)
    run.font.color.rgb = BLACK
    run.bold = bold


def add_heading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=True)
    return p


def add_body(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=False)
    return p


def add_bullet(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(f"- {text}")
    _set_run(run, bold=False)
    return p


def add_table(doc, header: list, rows: list):
    # No named table style is applied -- the source document's own tables
    # (checked before writing this) also have style=None; it doesn't define
    # a "Table Grid" style, so setting one here raises KeyError.
    table = doc.add_table(rows=1, cols=len(header))
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        _set_run(run, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            run = cells[i].paragraphs[0].add_run(str(val))
            _set_run(run, bold=False)
    return table


def build_week3_section(doc):
    add_heading(doc, "2.2. Week 3 — RL Baseline Vanilla: TD3 (Specific Objective 3)")
    add_body(doc,
        "Scope deviation, decided before any code was written this week: the "
        "original plan assigned Week 3 to the MPC/Gurobi baseline and RL to "
        "Weeks 4-5. This was deliberately reversed — Week 3 = RL vanilla "
        "(TD3), Week 4 = a perfect-information reference computed with a free "
        "solver (no Gurobi) plus PI-TD3, Week 5 = the full comparison. Reasons: "
        "RL training is the highest-lead-time, highest-risk item in the "
        "compressed 8-week plan, and Week 2 confirmed only a restricted, "
        "non-production Gurobi license is available in this environment, not "
        "an academic one — keeping MPC in Week 3 would have made the week "
        "depend on a licensing question that was never resolved.")

    add_heading(doc, "2.2.1. Training Configuration")
    add_body(doc,
        "Algorithm: TD3 (Twin Delayed DDPG) vanilla, via Stable-Baselines3. "
        "Chosen over SAC specifically because the physics-informed PI-TD3 "
        "method planned for Week 4 is defined as an ablation on top of TD3 "
        "in the source paper — using TD3 now keeps that comparison a "
        "single-factor ablation instead of two simultaneous changes.")
    add_body(doc,
        "Reward function: SqTrError_TrPenalty_UserIncentives (tracking error "
        "plus transformer-overload and user-dissatisfaction penalties), "
        "chosen over the environment's bare tracking-only default because it "
        "is the closer match to this thesis's declared success criteria. "
        "State function: PublicPST, the environment's own default for this "
        "problem variant. Both selected from what exists in the EV2Gym "
        "library, not newly written.")
    add_body(doc,
        "Every hyperparameter has a declared origin (SB3 default, TD3 paper, "
        "or set for this project's CPU budget) — network architecture "
        "reduced to [64, 64] from the paper's [400, 300], replay buffer "
        "reduced to 50,000 from SB3's default 1,000,000, both reduced "
        "specifically for the CPU-only budget. Training budget: an explicit "
        "5,000-timestep calibration run (18.09 steps/s measured) was "
        "presented as a table of 4 candidate budgets and the user confirmed "
        "60,000 timesteps per training seed (three seeds: 100, 101, 102) "
        "before any long run was launched — a drastic, declared reduction "
        "against the source papers' 5-48 hour HPC training runs.")

    add_heading(doc, "2.2.2. Results Obtained — TD3 vs. AFAP / Round Robin / Random Control")
    add_body(doc,
        "Training: all 3 seeds completed, 167.6 minutes (2.79 hours) total "
        "wall-clock, matching the calibration estimate almost exactly. "
        "Evaluated on the identical 50-cell grid (5 seeds x 10 days) already "
        "used for AFAP/Round Robin, plus a random-policy negative control "
        "(same action space, untrained, action-space seeded with the "
        "scenario seed).")
    add_table(doc,
        ["Algorithm", "EVs served", "Profit/Cost", "Avg. satisfaction", "Min. energy satisfaction", "Transformer overload (kWh)"],
        [
            ["AFAP (uncontrolled)", "13.44", "-45.73", "100%", "100.00", "5.33"],
            ["Round Robin", "13.44", "-44.51", "100%", "99.95", "0.00"],
            ["TD3 (seed 100)", "13.00", "-28.59", "99.6%", "93.06", "0.00"],
            ["TD3 (seed 101)", "10.40", "-36.69", "99.5%", "92.31", "0.00"],
            ["TD3 (seed 102)", "12.80", "-39.69", "98.4%", "86.71", "0.98"],
            ["Random policy (control)", "13.90", "-48.19", "100%", "100.00", "12.74"],
        ])

    add_heading(doc, "2.2.3. Interpretation")
    add_body(doc,
        "TD3 learns a real, demonstrated overload-avoidance behavior: it "
        "matches Round Robin's near-elimination of transformer overload "
        "across all 3 training seeds, and — critically — the "
        "random-policy control shows WORSE overload (12.74 kWh) than even "
        "unmanaged AFAP (5.33 kWh). Without that control, \"TD3 beats AFAP\" "
        "could not be distinguished from \"any policy that spreads power "
        "around beats AFAP at this station's 4:1 oversubscription ratio\" — "
        "the control rules that out.")
    add_body(doc,
        "This comes at a real, consistent cost: lower and seed-inconsistent "
        "EVs served (10.4-13.0 vs. 13.44 for both baselines), and materially "
        "lower worst-case user satisfaction (down to 86.7% for the worst "
        "seed, vs. ~100% for every non-RL policy tested). TD3 also tracks "
        "the power setpoint measurably WORSE than the much simpler Round "
        "Robin heuristic, despite optimizing a tracking-based reward — "
        "consistent with a declared reward-vs-evaluation-metric "
        "misalignment: the training reward and the reported evaluation "
        "metrics are related but not identical quantities. Training-seed "
        "dispersion is itself substantial and reported separately from "
        "scenario-level uncertainty (e.g. EVs served varies 21.5% and "
        "tracking error varies 48.9% purely from training-seed choice, "
        "holding the evaluation grid fixed).")

    add_heading(doc, "2.2.4. Limitations Identified")
    add_bullet(doc, "Training budget is a drastic, declared reduction against the source papers (2.79h total vs. 5-48h per single HPC run) — the observed weak learning signal must be read against this, not treated as a ceiling on what TD3 could achieve with more budget.")
    add_bullet(doc, "The apparently better profit figures for TD3 are not an intentional achievement — profitability has no representation in the training reward at all, and is very likely a side effect of serving fewer EVs.")
    add_bullet(doc, "Trained and evaluated on the single reference station only — the CPU budget did not extend to the station-size sensitivity sweep.")
    add_bullet(doc, "A perfect-information reference point does not exist yet, so no strategy tested so far (including TD3) can be described as the best achievable — only as the best-performing among the strategies actually tested.")

    add_heading(doc, "2.2.5. Next Steps (Week 4)")
    add_bullet(doc, "Implement the perfect-information reference (offline upper bound) using a free solver (HiGHS/CVXPY/OR-Tools) — no Gurobi dependency.")
    add_bullet(doc, "Add PI-TD3 as a physics-informed ablation on top of this week's vanilla-TD3 wrapper.")
    add_bullet(doc, "Begin the full cross-algorithm comparison scheduled for Week 5.")


def build_week4_section(doc):
    add_heading(doc, "2.3. Week 4 — Perfect-Information Oracle, Reward Falsification, and Reward Ablation (Specific Objective 3)")
    add_body(doc,
        "Scope amended twice, both kept in the project record rather than "
        "silently revised. Amendment 1 (2026-08-19): an academic Gurobi "
        "license became active on this machine (unlimited size, expires "
        "2027-08-14), so the perfect-information reference uses Gurobi "
        "directly rather than the free-solver fallback Week 2-3 had "
        "planned around — a reduction in deviation from the original plan, "
        "not a new one. Amendment 2 (2026-08-19): the originally planned "
        "trained \"PI-TD3\" arm was replaced with a falsified "
        "physics-adaptation investigation plus a reward ablation "
        "(TD3_TrackingOnly) that actually trained — no arm named PI_TD3 "
        "exists anywhere in this project's registry, code, or results.")

    add_heading(doc, "2.3.1. Perfect-Information Oracle")
    add_body(doc,
        "Two Gurobi-based oracle variants were built as a thin wrapper "
        "around EV2Gym's own PowerTrackingErrorrMin model (never edited in "
        "place): Optimal_Oracle_Tracking (tracking error only) and "
        "Optimal_Oracle_Balanced (tracking error plus a departure-"
        "satisfaction penalty, dimensionally corrected from a bug found in "
        "the library's own profit_max.py precedent). Both forced to G2V-only "
        "(matching every other algorithm compared in this thesis) and run "
        "over the same 50-cell evaluation grid (5 seeds x 10 days) used "
        "throughout this project. Result: both oracle variants dominate "
        "every online algorithm on tracking error by a wide margin (a "
        "2x-11x gap), with the two variants tying exactly on EVs served, "
        "transformer overload, and both satisfaction metrics but showing a "
        "small, fully explained, non-zero gap on tracking error and profit "
        "(an LP-tie-break vs. nonlinear-simulator model mismatch, "
        "quantified as an explicit noise floor rather than dismissed as "
        "unexplained variance).")

    add_heading(doc, "2.3.2. Physics-Informed Reward Adaptation: Falsified, Not Abandoned")
    add_body(doc,
        "The PI-TD3 paper's actual novel mechanism (read in full: "
        "arXiv:2510.12335v2) turned out to be a K-step differentiable-"
        "rollout actor update, not primarily a reward term — porting it "
        "faithfully would mean replacing Stable-Baselines3's training loop "
        "entirely, and this station has no power-flow model to "
        "differentiate through (simulate_grid=False). A thinner, reward-"
        "only adaptation was attempted instead, explicitly declared as "
        "such. Two structurally different designs were built and verified "
        "before either was trusted with training compute, and both failed, "
        "for two independent, well-diagnosed reasons: Design 1 (an "
        "instantaneous transformer-capacity margin) was rank-correlated "
        "1.0 with the reward's existing overload penalty at every weight "
        "tested — both are monotonic functions of the same instantaneous "
        "scalar, so no weight could ever make them disagree. Design 2 (a "
        "capacity-headroom term using the environment's own connected-"
        "fleet demand potential) broke that tie but was found to reward "
        "charging EVs to full faster regardless of whether that caused a "
        "real overload — confirmed by a heuristic that already nearly "
        "eliminates overload being penalized 32x more often than one that "
        "overloads routinely. The repair that would fix this needs each "
        "EV's departure time, which this project's chosen state function "
        "withholds deliberately, matching a real public station operator's "
        "actual information. Recorded as a conditional stretch goal for "
        "the Week 6 grid-enabled phase, not a commitment.")

    add_heading(doc, "2.3.3. TD3_TrackingOnly: A Reward Ablation")
    add_body(doc,
        "With the physics-adaptation line closed, the freed training "
        "budget went to a question Week 3 assumed rather than measured: "
        "does encoding transformer-overload and user-satisfaction "
        "penalties into the training reward actually buy anything over "
        "optimizing tracking error alone? TD3_TrackingOnly held every "
        "hyperparameter identical to Week 3's vanilla-TD3 arm (60,000 "
        "timesteps/seed, the same 3 training seeds, state function, "
        "network architecture) and changed only the reward function, to "
        "EV2Gym's own bare tracking-only default — a genuinely single-"
        "variable comparison.")
    add_table(doc,
        ["Algorithm", "Avg. satisfaction", "Min. energy satisfaction", "Transformer overload (kWh)", "Tracking error"],
        [
            ["Round Robin", "100.0%", "99.95", "0.00", "12,272.97"],
            ["TD3_vanilla (mean, 3 seeds)", "98.2%", "82.5", "1.36", "29,211.59"],
            ["TD3_TrackingOnly (mean, 3 seeds)", "98.6%", "86.4", "3.10", "32,286.50"],
            ["Optimal_Oracle_Tracking", "100.0%", "100.00", "0.00", "5,192.86"],
        ])

    add_heading(doc, "2.3.4. Results Obtained")
    add_body(doc,
        "Training: all 3 seeds completed, 168.43 minutes total wall-clock "
        "(within 1.4% of the pre-training calibration estimate). "
        "Evaluated on the identical 50-cell grid: 150 rows appended, "
        "bringing the main-grid registry to exactly 550 rows (11 "
        "algorithms x 50 cells). Pooled paired-bootstrap comparison "
        "(n=150, matched training-seed-to-training-seed) against "
        "TD3_vanilla: TD3_TrackingOnly achieves better profit (+5.0%) and "
        "satisfaction (+0.4% average, +7.6% worst-case), but worse "
        "tracking error (+13.0%), transformer overload (+1.75 kWh), and "
        "battery degradation (+2.6%) — every one of these differences is "
        "statistically significant (95% CI excludes zero). Against the "
        "oracle: every TD3 variant (both reward arms, all 6 checkpoints) "
        "sits farther from the tracking-error bound than the simple Round "
        "Robin heuristic does, including the arm trained solely to "
        "minimize tracking error.")

    add_heading(doc, "2.3.5. Interpretation")
    add_body(doc,
        "The reward ablation retroactively validates Week 3's reward "
        "choice rather than merely asserting it: the composite reward "
        "(tracking + overload + satisfaction penalties) achieves real, if "
        "modest and seed-dependent, improvements in tracking error AND "
        "overload avoidance over optimizing tracking alone — the two axes "
        "it was designed to improve — at a real cost in profit and "
        "worst-case satisfaction. The cross-cutting finding first observed "
        "in Week 3 is reinforced, not just repeated: every RL variant "
        "tested across both weeks tracks the power setpoint worse than the "
        "simple Round Robin heuristic, and this persists even for the arm "
        "trained exclusively on tracking error — ruling out reward-shaping "
        "distraction as the explanation. This is now the single most "
        "robust finding spanning this project's RL work to date.")

    add_heading(doc, "2.3.6. Limitations Identified")
    add_bullet(doc, "Both TD3 reward arms show substantial cross-training-seed dispersion (up to 63.8% relative spread on individual metrics) at the declared, reduced 60,000-timestep/seed budget — a real limitation of this project's compute scope, not resolved this week.")
    add_bullet(doc, "The proposal's declared \"<15% energy error\" success target has not yet been operationalized as a concrete percentage against a specific registry metric — flagged explicitly as an open item for Week 5's final comparison rather than resolved with a guessed formula.")
    add_bullet(doc, "The physics-informed reward line is closed for this station configuration (simulate_grid=False, PublicPST) — a faithful port remains possible only if Week 6-7's grid-enabled phase provides a real voltage/power-flow signal to adapt.")
    add_bullet(doc, "Voltage-band compliance (the proposal's third declared success target) remains untested — no voltage data exists until the grid-enabled phase.")

    add_heading(doc, "2.3.7. Next Steps (Week 5)")
    add_bullet(doc, "The full cross-algorithm comparison (all 11 arms: 2 heuristics, 1 random control, 6 TD3 checkpoints across 2 reward arms, 2 oracle variants) scheduled for Week 5.")
    add_bullet(doc, "Operationalize the proposal's \"<15% energy error\" target against a specific, named registry metric before the final comparison is written up.")
    add_bullet(doc, "Carry forward the reward-ablation finding (composite reward beats tracking-only, even on tracking error itself) as a design conclusion for any further RL work in Weeks 6-7.")


def build_week4_correction_section(doc):
    """Same-day acceptance review (2026-08-20), appended AFTER
    build_week4_section's own 2.3.1-2.3.7 -- not a rewrite of them.
    2.3.5/2.3.6/2.3.7 above (already saved in the real file before this
    review) overstated the reward-ablation verdict, understated the
    RL-vs-Round-Robin gap-to-oracle finding, and listed the energy-error
    target as still fully open even though a mapping has now been
    proposed. All three corrected here as a new, clearly-dated subsection,
    per this project's standing convention: never silently rewrite a
    prior claim, add a marked correction instead."""
    add_heading(doc, "2.3.8. Correction (2026-08-20, same-day acceptance review)")
    add_body(doc,
        "The results reported in 2.3.4 above are unchanged; three "
        "interpretive claims in 2.3.5-2.3.7 were reviewed the same day "
        "and found to be stated more strongly -- or more weakly -- than "
        "the evidence supports. Corrected here, not rewritten in place.")
    add_bullet(doc,
        "2.3.5 called the reward ablation a validation of Week 3's reward "
        "choice. Corrected: it is a trade-off, not a validation. "
        "TD3_vanilla wins on tracking error, energy tracking error, "
        "transformer overload, and battery degradation; TD3_TrackingOnly "
        "wins on profit and both satisfaction metrics -- two of this "
        "thesis's three declared objective axes (satisfaction, "
        "profitability, technical compliance). Neither dominates. The "
        "counterintuitive direction -- optimizing tracking error alone "
        "producing a worse realized tracking error -- is recorded as a "
        "hypothesis (reduced training-budget reward shaping aiding "
        "credit assignment), not an established mechanism; a longer "
        "single-seed run would test it and is explicitly out of scope "
        "this week.")
    add_bullet(doc,
        "2.3.5 mentioned the RL-vs-Round-Robin tracking-error gap only in "
        "passing. Corrected: this is the week's real headline and is "
        "stated as such. Under a 60,000-timestep CPU training budget, "
        "reinforcement learning does not beat the simple Round Robin "
        "heuristic on tracking error, by a wide margin, consistently "
        "across two reward functions and six independent training seeds "
        "(Round Robin: 136.3% gap to the tracking-error oracle bound; "
        "every TD3 checkpoint: 422-612%). Bounded explicitly to this "
        "training budget, not claimed as a general property of RL vs. "
        "heuristics. Consequence for Objective 4's recommended-strategy "
        "question, surfaced now rather than left for Week 6: on current "
        "evidence the recommendation points toward Round Robin, not an "
        "RL method -- a legitimate, defensible conclusion this thesis is "
        "fully able to reach.")
    add_bullet(doc,
        "2.3.6/2.3.7 listed the \"<15% energy error\" target as fully "
        "open. Corrected: a concrete mapping has now been proposed (not "
        "yet adopted) -- (1 - average_user_satisfaction) x 100, using "
        "EV2Gym's own per-EV satisfaction ratio against each EV's "
        "requested (not idealized-max-rate) energy, read from source "
        "(ev2gym/models/ev.py, ev2gym/utilities/utils.py). Computed for "
        "every algorithm and arm tested through Week 4: worst observed "
        "value is 2.57% (TD3_vanilla, cross-check definition), over 5x "
        "under the 15% threshold -- every arm clears the target "
        "comfortably under either of two candidate definitions checked "
        "against each other. Week 5's task is to adopt or revise this "
        "proposal, not to define one from scratch.")


if __name__ == "__main__":
    if not os.path.exists(PROGRESS_LOG_PATH):
        raise FileNotFoundError(
            f"{PROGRESS_LOG_PATH!r} not found relative to the current working "
            f"directory. Run this script with PYTHONPATH=. from the repo root "
            f"(PG_IIND/), so the relative path resolves to the parent 'PG "
            f"Industrial' folder where this file actually lives."
        )
    doc = Document(PROGRESS_LOG_PATH)
    n_paragraphs_before = len(doc.paragraphs)
    # Week 3's section was already appended to the real file on 2026-08-13,
    # and Week 4's own build_week4_section was already appended earlier
    # today (2026-08-20) -- the file already contains "2.2. Week 3" and
    # "2.3. Week 4" headings. Calling either build_weekN_section again here
    # would duplicate it. Only the same-day correction runs now, as its own
    # new subsection (2.3.8) -- see build_week4_correction_section's own
    # docstring for why this is a new subsection, not an in-place edit.
    build_week4_correction_section(doc)
    doc.save(PROGRESS_LOG_PATH)
    print(f"Appended Week 4 correction section to {PROGRESS_LOG_PATH} "
          f"({n_paragraphs_before} paragraphs before -> {len(doc.paragraphs)} after). "
          f"Nothing before paragraph {n_paragraphs_before} was modified.")
